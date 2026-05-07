"""Generate Word document with current bot trading rules from config.py."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime, timezone

from config import (
    BANKROLL, YES_ENTRY_STEPS, YES_EXIT_STEPS, NO_ENTRY_STEPS,
    WTI_SELL_NO_BELOW, YES_PROFIT_TARGET,
    STOP_CEASEFIRE_WTI, STOP_NO_WTI, STOP_THETA_DAYS, STOP_THETA_WTI_MOVE,
    STOP_PORTFOLIO_FLOOR, DEAD_ZONE_START_ET, DEAD_ZONE_END_ET,
    DEADLINE_DATE, SETTLEMENT_TRIGGER, THETA_SCALING,
    DIVERGENCE_RATIO_NORMAL, DIVERGENCE_ALERT_MULT,
    DIVERGENCE_WINDOW_HOURS, DIVERGENCE_REDUCE_PCT,
    MIN_BET_USD, SCAN_INTERVAL, ORDER_TTL_SECONDS, MAX_SLIPPAGE,
)

BOT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_PATH = os.path.join(BOT_DIR, "Oil_Bot_Rules.docx")


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def generate():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('Oil Swing Trading Bot — Trading Rules', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Polymarket: "Will Crude Oil (CL) hit $100 by end of March?"')
    run.italic = True
    run.font.size = Pt(12)

    now_msk = datetime.now(timezone.utc).strftime("%d.%m.%Y %H:%M UTC")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"Generated: {now_msk}")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(128, 128, 128)

    # === OVERVIEW ===
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        f"Bankroll: ${BANKROLL:.0f}\n"
        f"Strategy: Swing trade YES and NO tokens based on WTI crude oil price.\n"
        f"Buy YES when WTI drops (oil cheap = YES cheap).\n"
        f"Buy NO when WTI rises (oil expensive = NO cheap).\n"
        f"Deadline: {DEADLINE_DATE} (close all positions).\n"
        f"Scan interval: every {SCAN_INTERVAL} seconds."
    )

    # === YES ENTRY ===
    doc.add_heading('2. Entry YES — Gradient (3 steps as WTI drops)', level=1)
    doc.add_paragraph("Buy YES in 3 steps with increasing size as WTI falls further:")

    yes_entry_rows = []
    for i, (trigger, bet) in enumerate(YES_ENTRY_STEPS):
        labels = ["Scout (разведка)", "Confirm (подтверждение)", "Main position (основная)"]
        yes_entry_rows.append([
            f"Step {i+1}",
            f"WTI <= ${trigger:.0f}",
            f"${bet:.0f}",
            labels[i] if i < len(labels) else ""
        ])
    add_table(doc, ["Step", "WTI Trigger", "Bet Size", "Purpose"], yes_entry_rows)

    doc.add_paragraph(
        f"\nTotal max YES investment: ${sum(b for _, b in YES_ENTRY_STEPS):.0f}"
    )

    # === YES EXIT ===
    doc.add_heading('3. Exit YES — Two conditions (OR logic)', level=1)

    doc.add_heading('3a. Profit-based exit', level=2)
    doc.add_paragraph(
        f"Sell YES if position profit >= {YES_PROFIT_TARGET:.0%} from entry price.\n"
        f"Applies to any YES position regardless of WTI level."
    )

    doc.add_heading('3b. WTI-based gradient exit (3 steps)', level=2)
    doc.add_paragraph("Sell YES in 3 steps as WTI rises:")

    yes_exit_rows = []
    for i, (trigger, pct) in enumerate(YES_EXIT_STEPS):
        yes_exit_rows.append([
            f"Step {i+1}",
            f"WTI >= ${trigger:.0f}",
            f"{pct:.0%}",
        ])
    add_table(doc, ["Step", "WTI Trigger", "Sell % of position"], yes_exit_rows)

    # === NO ENTRY ===
    doc.add_heading('4. Entry NO — Gradient (3 steps as WTI rises)', level=1)
    doc.add_paragraph("Buy NO in 3 steps with per-step profit targets:")

    no_entry_rows = []
    for i, (trigger, bet, target) in enumerate(NO_ENTRY_STEPS):
        no_entry_rows.append([
            f"Step {i+1}",
            f"WTI >= ${trigger:.0f}",
            f"${bet:.0f}",
            f"+{target:.0%}",
        ])
    add_table(doc, ["Step", "WTI Trigger", "Bet Size", "Profit Target"], no_entry_rows)

    doc.add_paragraph(
        f"\nTotal max NO investment: ${sum(b for _, b, _ in NO_ENTRY_STEPS):.0f}"
    )

    # === NO EXIT ===
    doc.add_heading('5. Exit NO — Two conditions (OR logic)', level=1)

    doc.add_heading('5a. Profit-based exit (per-step targets)', level=2)
    doc.add_paragraph(
        "Each NO position has its own profit target based on entry step:\n"
        "- Step 1 (WTI $97): sell at +50% profit\n"
        "- Step 2 (WTI $98): sell at +100% profit\n"
        "- Step 3 (WTI $99): sell at +150% profit\n\n"
        "Logic: cheaper entry (higher WTI) = higher profit target,\n"
        "because NO bought at extreme WTI levels has more upside potential."
    )

    doc.add_heading('5b. WTI-based exit', level=2)
    doc.add_paragraph(
        f"Sell all NO when WTI <= ${WTI_SELL_NO_BELOW:.0f}\n"
        f"(Oil dropped back — NO loses value, exit to lock in remaining value)"
    )

    # === STOP-LOSSES ===
    doc.add_heading('6. Stop-Loss Rules', level=1)

    stop_rows = [
        ["Ceasefire", f"WTI < ${STOP_CEASEFIRE_WTI:.0f}", "Sell ALL YES",
         "Oil collapsed, YES worthless"],
        ["NO Risk", f"WTI > ${STOP_NO_WTI:.0f}", "Sell ALL NO",
         "Oil near $100, NO at risk of resolution"],
        ["Theta Decay", f"YES held > {STOP_THETA_DAYS} days,\nWTI moved < ${STOP_THETA_WTI_MOVE:.0f}",
         "Sell that YES position", "Time decay eating profits"],
        ["Portfolio Floor", f"Portfolio < ${STOP_PORTFOLIO_FLOOR:.0f}", "Sell ALL, stop bot",
         "Capital preservation"],
        ["Settlement", f"WTI settlement >= ${SETTLEMENT_TRIGGER:.0f}", "Market resolves YES",
         "Only CME settlement price counts"],
    ]
    add_table(doc,
              ["Stop", "Condition", "Action", "Reason"],
              stop_rows,
              col_widths=[1.0, 1.5, 1.3, 1.7])

    # === TIME RULES ===
    doc.add_heading('7. Time Rules', level=1)

    et_start = DEAD_ZONE_START_ET
    et_end = DEAD_ZONE_END_ET
    msk_start = (et_start + 7) % 24
    msk_end = (et_end + 7) % 24

    doc.add_paragraph(
        f"Dead zone: {et_start}:00-{et_end}:00 ET "
        f"({msk_start}:00-{msk_end}:00 MSK)\n"
        f"No new positions during this window (CME settlement at 14:30 ET / 21:30 MSK).\n\n"
        f"Buy window: after {et_end}:00 ET ({msk_end}:00 MSK)\n"
        f"Sell window: before {et_start}:00 ET ({msk_start}:00 MSK)\n\n"
        f"Deadline: {DEADLINE_DATE} — close all positions and stop bot.\n"
        f"(3 trading days before March 31 contract expiry)"
    )

    # === THETA SCALING ===
    doc.add_heading('8. Theta Scaling (Position Size Reduction)', level=1)
    doc.add_paragraph("As deadline approaches, reduce new position sizes:")

    theta_rows = []
    for min_days, mult in THETA_SCALING:
        if min_days == 0:
            label = "0-2 days"
        else:
            label = f"{min_days}+ days"
        theta_rows.append([label, f"{mult:.0%}"])
    add_table(doc, ["Days Remaining", "Position Size"], theta_rows)

    # === DIVERGENCE ===
    doc.add_heading('9. Divergence Detection', level=1)
    doc.add_paragraph(
        f"Normal ratio: $1 WTI change = ~{DIVERGENCE_RATIO_NORMAL}pp YES change.\n"
        f"Alert threshold: actual move > {DIVERGENCE_ALERT_MULT}x expected.\n"
        f"Measurement window: {DIVERGENCE_WINDOW_HOURS} hours.\n\n"
        f"During divergence:\n"
        f"- Reduce new position sizes to {DIVERGENCE_REDUCE_PCT:.0%} for 2 hours.\n"
        f"- Flash crash protection: if YES drops >15pp but WTI drops <$2 — do NOT sell.\n"
    )

    # === ORDER EXECUTION ===
    doc.add_heading('10. Order Execution', level=1)
    doc.add_paragraph(
        f"Order type: Limit (GTC)\n"
        f"Order timeout: {ORDER_TTL_SECONDS} sec ({ORDER_TTL_SECONDS // 60} min) — cancel if not filled.\n"
        f"Max slippage: {MAX_SLIPPAGE:.0%} from mid price.\n"
        f"Min bet: ${MIN_BET_USD:.0f}\n"
        f"Min shares: 5 (Polymarket minimum)\n"
        f"Failed sell retry: 3% lower price."
    )

    doc.save(OUTPUT_PATH)
    print(f"Word document saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    generate()
